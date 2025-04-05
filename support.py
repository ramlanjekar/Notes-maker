from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

def convert_md_to_docx(md_path, output_path):
    doc = Document()
    
    # Configure heading styles
    heading_styles = {
        1: {"name": "Heading 1", "font_size": 16, "bold": True},
        2: {"name": "Heading 2", "font_size": 14, "bold": True},
        3: {"name": "Heading 3", "font_size": 12, "bold": True}
    }
    
    # Configure bullet styles (level: symbol, font_size)
    bullet_config = {
        1: {"symbol": "•", "font_size": 10},
        2: {"symbol": "➢", "font_size": 10},
        3: {"symbol": "-", "font_size": 10},
        4: {"symbol": "–", "font_size": 10},
        5: {"symbol": "⁃", "font_size": 10}
    }
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.readlines()
    
    current_list_level = 0
    for line in content:
        line = line.rstrip()
        
        # Pre-clean any stray ** that aren't for formatting
        line = clean_stray_asterisks(line)
        
        # Headings
        if line.startswith('# '):
            add_heading(doc, line[2:], 1, heading_styles)
        elif line.startswith('## '):
            add_heading(doc, line[3:], 2, heading_styles)
        elif line.startswith('### '):
            add_heading(doc, line[4:], 3, heading_styles)
        
        # Bullet Lists
        elif line.startswith('* ') or re.match(r'^ {1,8}\* ', line) or line.startswith('- ') or re.match(r'^ {1,8}- ', line):
            process_list_item(doc, line, bullet_config)
        
        # Plain text
        elif line.strip():  # Only process non-empty lines
            add_paragraph(doc, line)
    
    doc.save(output_path)
    print(f"Lecture docx is prepared in {output_path}\n")

def clean_stray_asterisks(line):
    """Clean different patterns of stray ** in the text"""
    # Replace patterns like "word:** text" with "word: text"
    line = re.sub(r':\*\*(\s)', r':\1', line)
    
    # Replace pattern like "word:** at the end of line
    line = re.sub(r':\*\*$', r':', line)
    
    # Replace :**" with :"
    line = re.sub(r':\*\*"', r':"', line)
    
    # Replace standalone ** at end of line
    line = re.sub(r'\*\*$', '', line)
    
    # Replace patterns like "word:**word" with "word:word"
    line = re.sub(r':\*\*([A-Za-z])', r':\1', line)
    
    # Identify and keep legitimate bold markup
    line = re.sub(r'\*\*(.*?):\*\*', r'**\1:**', line)
    
    return line
    
def add_heading(doc, text, level, styles):
    """Add styled heading with decreasing font size"""
    # Clean any stray asterisks
    text = clean_stray_asterisks(text)
    
    heading = doc.add_heading(text, level=level)
    apply_font_formatting(heading, styles[level]["font_size"], bold=styles[level]["bold"])

def process_list_item(doc, line, config):
    """Process nested lists with proper indentation"""
    indent = len(re.match(r'^ *', line).group(0))
    level = (indent // 4) + 1  # 4 spaces per indent level
    
    # Get formatting from config
    fmt = config.get(min(level, 5))
    if fmt is None:
        fmt = config[5]
    
    p = doc.add_paragraph()
    
    # Apply indentation before the bullet point
    p.paragraph_format.left_indent = Cm(0.5 * level)  # Indentation before bullet
    p.paragraph_format.first_line_indent = Cm(-0.25)  # Hanging indent for bullet
    
    # Add bullet symbol
    bullet_run = p.add_run(fmt["symbol"] + " ")  # Changed tab to space
    bullet_run.font.size = Pt(fmt["font_size"])
    
    # Process text with markdown formatting - clean line of * or - prefix first
    clean_text = line.lstrip('* -').lstrip()
    
    # Clean any stray asterisks
    clean_text = clean_stray_asterisks(clean_text)
    
    process_markdown(p, clean_text, fmt["font_size"])

def add_paragraph(doc, text):
    """Add regular paragraph with markdown formatting"""
    if not text.strip():
        doc.add_paragraph()  # Add empty paragraph for blank lines
        return
    
    # Clean any stray asterisks
    text = clean_stray_asterisks(text)
    
    p = doc.add_paragraph()
    process_markdown(p, text, 11)  # Default size 11pt

def process_markdown(p, text, base_size):
    """Handle markdown formatting within text"""
    # First, identify legitimate bold text sections (text between paired **)
    bold_pattern = re.compile(r'\*\*(.*?)\*\*')
    last_end = 0
    parts = []
    
    for match in bold_pattern.finditer(text):
        start, end = match.span()
        
        # Add text before this match
        if start > last_end:
            parts.append((text[last_end:start], False))
        
        # Add the bold text (without the **)
        parts.append((match.group(1), True))
        
        last_end = end
    
    # Add any remaining text
    if last_end < len(text):
        remaining_text = text[last_end:]
        # Clean any remaining stray ** that weren't part of a pair
        remaining_text = re.sub(r'\*\*', '', remaining_text)
        parts.append((remaining_text, False))
    
    # If no parts were found, just use the whole text (with any ** removed)
    if not parts:
        text = re.sub(r'\*\*', '', text)
        parts = [(text, False)]
    
    # Add all parts with appropriate formatting
    for text_part, is_bold in parts:
        run = p.add_run(text_part)
        run.bold = is_bold
        run.font.size = Pt(base_size)

def apply_font_formatting(paragraph, size, bold=False):
    """Apply consistent formatting to paragraph"""
    for run in paragraph.runs:
        run.font.size = Pt(size)
        run.bold = bold

# Usage
# convert_md_to_docx('lecture_notes.md', 'output/formatted_output10.docx')